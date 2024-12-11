import os
import sys
import sqlite3

from docx import Document
from PyQt5 import QtWidgets, QtCore, QtGui
from PyQt5.QtWidgets import QLabel, QMessageBox
from PyQt5.QtGui import QFont

from MainWindow import Ui_MainWindow  # Импорт интерфейса основного окна
from Login import Ui_Login  # Импорт интерфейса логина


class Login(QtWidgets.QDialog):
    def __init__(self):
        super().__init__()
        self.ui = Ui_Login()  # Подключаем UI логина
        self.ui.setupUi(self)  # Настраиваем UI для QDialog

        # Установим режим отображения пароля
        self.ui.lineEdit_2.setEchoMode(QtWidgets.QLineEdit.Password)  # Скрывать текст
        # Подключаем событие для кнопки "Войти"
        self.ui.pushButton.clicked.connect(self.handle_login)

        # Создаем QLabel для отображения сообщений об ошибках
        self.error_label = QLabel(self)
        self.error_label.setStyleSheet("color: red;")  # Устанавливаем красный цвет текста
        self.error_label.setVisible(False)  # Сначала скрываем лейбл

        # Устанавливаем шрифт и выравнивание
        font = QFont("Segoe UI Semibold", 8, QFont.Bold)
        self.error_label.setFont(font)
        self.error_label.setAlignment(QtCore.Qt.AlignCenter)  # Выравнивание по центру

        # Добавляем error_label в вертикальный макет под полями ввода
        self.ui.verticalLayout_7.addWidget(self.error_label)

    def handle_login(self):
        """Обработка логина"""
        username = self.ui.lineEdit.text()
        password = self.ui.lineEdit_2.text()

        # Проверка логина и пароля
        if username == "admin" and password == "1234":
            self.accept()  # Закрываем окно логина и возвращаем True
        else:
            # Отображение сообщения об ошибке
            self.error_label.setText("Неправильный логин или пароль")
            self.error_label.setVisible(True)  # Показываем QLabel


class MyApp(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.ui = Ui_MainWindow()  # Подключаем UI основного окна
        self.ui.setupUi(self)  # Настраиваем UI для QMainWindow

        self.create_database()
        # Разрешаем изменение размера окна и показываем кнопки сворачивания/развертывания
        self.setWindowFlags(self.windowFlags() | QtCore.Qt.WindowMinMaxButtonsHint | QtCore.Qt.WindowCloseButtonHint)

        # Скрываем все виджеты по умолчанию
        self.hideAllPages()

        # Переключение между страницами с помощью QStackedWidget
        self.ui.stackedWidget.addWidget(self.ui.ReportpagePatients)  # Добавляем страницу пациентов
        self.ui.stackedWidget.addWidget(self.ui.ReportpageDoc)  # Добавляем страницу врачей
        self.ui.stackedWidget.addWidget(self.ui.ReportpageAppointment)  # Добавляем страницу записей

        # Соединение кнопок с соответствующими методами
        self.ui.Table1.clicked.connect(self.show_patients_table)
        self.ui.Table2.clicked.connect(self.show_doctors_table)
        self.ui.Table3.clicked.connect(self.show_appointments_table)
        self.ui.Tab1.clicked.connect(self.show_patients_table)
        self.ui.Tab2.clicked.connect(self.show_doctors_table)
        self.ui.Tab3.clicked.connect(self.show_appointments_table)
        self.ui.Bt1.clicked.connect(self.show_patients_table)
        self.ui.Bt2.clicked.connect(self.show_doctors_table)
        self.ui.Bt3.clicked.connect(self.show_appointments_table)
        self.ui.Lab.clicked.connect(lambda: self.showWidget('lab'))
        self.ui.Report.clicked.connect(lambda: self.showWidget('report'))
        self.ui.Statistics.clicked.connect(lambda: self.showWidget('statistics'))
        self.ui.Settings.clicked.connect(lambda: self.showWidget('settings'))
        self.ui.Help.clicked.connect(lambda: self.showWidget('help'))
        self.ui.Exit.clicked.connect(self.exit_to_login)  # Подключаем событие выхода
        self.ui.AddPatients.clicked.connect(self.add_patient)
        self.ui.AddDoctors.clicked.connect(self.add_doctor)
        self.ui.Delete.clicked.connect(self.delete_patient)
        self.ui.Delete2.clicked.connect(self.delete_doctor)
        self.ui.Delete3.clicked.connect(self.delete_appointment)
        self.ui.Edit.clicked.connect(self.show_right_menu)
        self.ui.Edit2.clicked.connect(self.show_right_menu)
        self.ui.Edit3.clicked.connect(self.add_appointment)
        self.ui.Save.clicked.connect(self.save_appointments)
        self.ui.Patients.clicked.connect(self.show_right_menu)
        self.ui.Doctors.clicked.connect(self.show_right_menu_doctors)
        self.ui.Patients2.clicked.connect(self.show_right_menu)
        self.ui.Doctors2.clicked.connect(self.show_right_menu_doctors)
        self.ui.Dismiss.clicked.connect(self.hide_right_menu)
        self.ui.Dismiss2.clicked.connect(self.hide_right_menu_doctors)
        self.ui.Menu.clicked.connect(self.hideShowLeftBar)
        self.ui.searchLineEdit.returnPressed.connect(self.highlight_patient)

        self.ui.stackedWidget.setCurrentWidget(self.ui.Statisticpage)

        self.load_patients()
        self.load_doctors()
        self.load_appointments()
        self.load_patient_statistics()

    def create_database(self):
        """Создание базы данных и таблиц"""
        connection = sqlite3.connect('db/hospital_management.db')
        cursor = connection.cursor()

        # Создание таблицы для пациентов
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS Patients (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            full_name TEXT NOT NULL,
            birthday TEXT,
            address TEXT,
            gender TEXT,
            phone_number TEXT
        )
        ''')

        # Создание таблицы для врачей
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS Doctors (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            full_name TEXT NOT NULL,
            job_title TEXT
        )
        ''')

        # Создание таблицы для записей
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS Appointments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            patient_id INTEGER,
            doctor_id INTEGER,
            appointment_date TEXT,
            date_of_application TEXT,  -- Новый столбец для даты обращения
            status TEXT,                -- Новый столбец для статуса посещения
            FOREIGN KEY (patient_id) REFERENCES Patients (id),
            FOREIGN KEY (doctor_id) REFERENCES Doctors (id)
        )
        ''')

        connection.commit()  # Сохраняем изменения
        connection.close()  # Закрываем соединение

    def add_patient(self):
        # Получаем данные из полей ввода
        full_name = self.ui.FIO.text().strip()
        birthday = self.ui.Birthday.text().strip()
        address = self.ui.Adres.text().strip()
        gender = 'М' if self.ui.Man.isChecked() else 'Ж'
        phone_number = self.ui.Phone.text().strip()

        if full_name:  # Проверка на пустое имя
            connection = sqlite3.connect('db/hospital_management.db')
            cursor = connection.cursor()
            cursor.execute(
                "INSERT INTO Patients (full_name, birthday, address, gender, phone_number) VALUES (?, ?, ?, ?, ?)",
                (full_name, birthday, address, gender, phone_number))
            connection.commit()
            connection.close()

            self.load_patients()  # Обновляем таблицу с пациентами
        else:
            QMessageBox.warning(self, "Ошибка", "Имя пациента не может быть пустым.")

    def delete_patient(self):
        row_index = self.ui.tableWidgetPatients.currentRow()
        if row_index >= 0:
            patient_id = self.ui.tableWidgetPatients.item(row_index, 0).text()

            # Создаем окно подтверждения
            reply = QMessageBox()
            reply.setIcon(QMessageBox.Question)
            reply.setText("Вы действительно хотите удалить пациента?")
            reply.setWindowTitle("Подтверждение")
            reply.setStandardButtons(QMessageBox.Yes | QMessageBox.No)

            # Устанавливаем текст кнопок на русском
            reply.button(QMessageBox.Yes).setText("Да")
            reply.button(QMessageBox.No).setText("Нет")

            # Показываем окно и получаем ответ
            if reply.exec_() == QMessageBox.Yes:
                connection = sqlite3.connect('db/hospital_management.db')
                cursor = connection.cursor()
                cursor.execute("DELETE FROM Patients WHERE id = ?", (patient_id,))
                connection.commit()
                connection.close()

                self.load_patients()  # Обновляем таблицу
        else:
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите пациента для удаления.")

    def edit_patient(self):
        row_index = self.ui.tableWidgetPatients.currentRow()
        if row_index >= 0:
            patient_id = self.ui.tableWidgetPatients.item(row_index, 0).text()  # Получаем ID пациента
            full_name = self.ui.FIO.text()
            birthday = self.ui.Birthday.text()
            address = self.ui.Adres.text()
            gender = 'М' if self.ui.Man.isChecked() else 'Ж'
            phone_number = self.ui.Phone.text()
            patient_data = self.get_patient_by_id(patient_id)  # Получаем данные пациента

            connection = sqlite3.connect('db/hospital_management.db')
            cursor = connection.cursor()
            cursor.execute("""
                UPDATE Patients 
                SET full_name = ?, birthday = ?, address = ?, gender = ?, phone_number = ?
                WHERE id = ?
            """, (full_name, birthday, address,gender, phone_number, patient_id))
            connection.commit()
            connection.close()

            self.load_patients()  # Обновляем таблицу

            if patient_data:
                # Обновляем поля ввода на основании полученных данных
                self.ui.FIO.setText(patient_data[1])
                self.ui.Birthday.setText(patient_data[2])
                self.ui.Adres.setText(patient_data[3])
                self.ui.Man.setChecked(patient_data[4] == 'М')
                self.ui.Woman.setChecked(patient_data[4] == 'Ж')
                self.ui.Phone.setText(patient_data[5])
            else:
                QMessageBox.warning(self, "Ошибка", "Пациент не найден.")
        else:
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите пациента для редактирования.")

    def load_patients(self):
        connection = sqlite3.connect('db/hospital_management.db')
        cursor = connection.cursor()
        cursor.execute("SELECT * FROM Patients")
        self.ui.tableWidgetPatients.setRowCount(0)
        for row in cursor.fetchall():
            row_position = self.ui.tableWidgetPatients.rowCount()
            self.ui.tableWidgetPatients.insertRow(row_position)
            for column, data in enumerate(row):
                item = QtWidgets.QTableWidgetItem(str(data))
                self.ui.tableWidgetPatients.setItem(row_position, column, item)

                # Подключаем обработчик изменения ячейки
                if column == 1:  # Изменение ФИО (например, или другой важный столбец)
                    item.setFlags(QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEnabled)

        # Подключаем сигнал itemChanged ко всему виджету
        self.ui.tableWidgetPatients.itemChanged.connect(
            lambda item: self.update_patient_data(item.row(), item.column()))
        connection.close()

    def add_doctor(self):
        full_name = self.ui.FIODoctors.text()
        job_title = self.ui.Job.text()

        connection = sqlite3.connect('db/hospital_management.db')
        cursor = connection.cursor()
        cursor.execute("INSERT INTO Doctors (full_name, job_title) VALUES (?, ?)",
                       (full_name, job_title))
        connection.commit()
        connection.close()

        self.load_doctors()  # Обновляем таблицу врачей

    def delete_doctor(self):
        row_index = self.ui.tableWidgetDoc.currentRow()
        if row_index >= 0:
            doctor_id = self.ui.tableWidgetDoc.item(row_index, 0).text()

            # Создаем окно подтверждения
            reply = QMessageBox()
            reply.setIcon(QMessageBox.Question)
            reply.setText("Вы действительно хотите удалить врача?")
            reply.setWindowTitle("Подтверждение")
            reply.setStandardButtons(QMessageBox.Yes | QMessageBox.No)

            # Устанавливаем текст кнопок на русском
            reply.button(QMessageBox.Yes).setText("Да")
            reply.button(QMessageBox.No).setText("Нет")

            # Показываем окно и получаем ответ
            if reply.exec_() == QMessageBox.Yes:
                connection = sqlite3.connect('db/hospital_management.db')
                cursor = connection.cursor()
                cursor.execute("DELETE FROM Doctors WHERE id = ?", (doctor_id,))
                connection.commit()
                connection.close()

                self.load_doctors()  # Обновляем таблицу врачей
        else:
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите врача для удаления.")

    def edit_doctor(self):
        row_index = self.ui.tableWidgetDoc.currentRow()
        if row_index >= 0:
            doctor_id = self.ui.tableWidgetDoc.item(row_index, 0).text()  # Получаем ID врача
            full_name = self.ui.FIODoctors.text()
            job_title = self.ui.Job.text()

            connection = sqlite3.connect('db/hospital_management.db')
            cursor = connection.cursor()
            cursor.execute("""
                UPDATE Doctors 
                SET full_name = ?, job_title = ?
                WHERE id = ?
            """, (full_name, job_title, doctor_id))
            connection.commit()
            connection.close()

            self.load_doctors()  # Обновляем таблицу
        else:QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите врача для редактирования.")

    def load_doctors(self):
        connection = sqlite3.connect('db/hospital_management.db')
        cursor = connection.cursor()
        cursor.execute("SELECT * FROM Doctors")
        self.ui.tableWidgetDoc.setRowCount(0)
        for row in cursor.fetchall():
            row_position = self.ui.tableWidgetDoc.rowCount()
            self.ui.tableWidgetDoc.insertRow(row_position)
            for column, data in enumerate(row):
                item = QtWidgets.QTableWidgetItem(str(data))
                self.ui.tableWidgetDoc.setItem(row_position, column, item)

        # Подключаем сигнал itemChanged для обновления информации о врачах
        self.ui.tableWidgetDoc.itemChanged.connect(lambda item: self.update_doctor_data(item.row(), item.column()))
        connection.close()

    def add_appointment(self):
        """Добавляет новую запись в таблицу записей."""
        patient_row = self.ui.tableWidgetPatients.currentRow()
        doctor_row = self.ui.tableWidgetDoc.currentRow()

        if patient_row >= 0 and doctor_row >= 0:  # Проверяем, выбраны ли оба
            patient_full_name = self.ui.tableWidgetPatients.item(patient_row, 1).text()  # Получаем ФИО пациента
            doctor_full_name = self.ui.tableWidgetDoc.item(doctor_row, 1).text()  # Получаем ФИО врача

            # Вставляем новую строку в tableWidgetAppointment
            row_position = self.ui.tableWidgetAppointment.rowCount()
            self.ui.tableWidgetAppointment.insertRow(row_position)

            # Генерируем ID записи
            appointment_id = self.get_next_appointment_id()
            self.ui.tableWidgetAppointment.setItem(row_position, 0,
                                                   QtWidgets.QTableWidgetItem(str(appointment_id)))  # ID записи
            # Вставляем ФИО пациента и врача
            self.ui.tableWidgetAppointment.setItem(row_position, 1,
                                                   QtWidgets.QTableWidgetItem(patient_full_name))  # ФИО пациента
            self.ui.tableWidgetAppointment.setItem(row_position, 2,
                                                   QtWidgets.QTableWidgetItem(doctor_full_name))  # ФИО врача

            # Создаем ячейку для даты приема и делаем ее редактируемой
            appointment_date_item = QtWidgets.QTableWidgetItem()
            appointment_date_item.setFlags(
                QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEnabled)
            self.ui.tableWidgetAppointment.setItem(row_position, 3, appointment_date_item)  # Дата приема

            # Создаем ячейку для даты обращения и делаем ее редактируемой
            date_application_item = QtWidgets.QTableWidgetItem()
            date_application_item.setFlags(
                QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEnabled)
            self.ui.tableWidgetAppointment.setItem(row_position, 4, date_application_item)  # Дата обращения

            # Создаем ячейку для статуса посещения и оставляем ее пустой, чтобы пользователь мог вводить
            status_item = QtWidgets.QTableWidgetItem("")  # Пустая ячейка для статуса посещения
            status_item.setFlags(
                QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEnabled)
            self.ui.tableWidgetAppointment.setItem(row_position, 5, status_item)  # Статус посещения

            # Уведомление пользователя
            QMessageBox.information(self, "Информация", "Запись добавлена.")
        else:
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите пациента и врача для создания записи.")

    def get_next_appointment_id(self):
        """Возвращает следующий доступный ID для новой записи приема."""
        connection = sqlite3.connect('db/hospital_management.db')
        cursor = connection.cursor()
        cursor.execute("SELECT MAX(id) FROM Appointments")
        max_id = cursor.fetchone()[0]  # Получаем максимальный ID из таблицы
        connection.close()
        return (max_id + 1) if max_id is not None else 1  # Если нет записей, начинаем с 1

    def delete_appointment(self):
        """Удаляет выбранную запись из таблицы."""
        row_index = self.ui.tableWidgetAppointment.currentRow()
        if row_index >= 0:
            appointment_id = self.ui.tableWidgetAppointment.item(row_index, 0).text()

            # Создаем окно подтверждения
            reply = QMessageBox()
            reply.setIcon(QMessageBox.Question)
            reply.setText("Вы действительно хотите удалить эту запись?")
            reply.setWindowTitle("Подтверждение")
            reply.setStandardButtons(QMessageBox.Yes | QMessageBox.No)

            # Устанавливаем текст кнопок на русском
            reply.button(QMessageBox.Yes).setText("Да")
            reply.button(QMessageBox.No).setText("Нет")

            # Показываем окно и получаем ответ
            if reply.exec_() == QMessageBox.Yes:
                connection = sqlite3.connect('db/hospital_management.db')
                cursor = connection.cursor()
                cursor.execute("DELETE FROM Appointments WHERE id = ?", (appointment_id,))
                connection.commit()
                connection.close()

                self.load_appointments()  # Обновляем таблицу записей
        else:
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите запись для удаления.")

    def edit_appointment(self):
        """Редактирует выбранную запись в таблице."""
        row_index = self.ui.tableWidgetAppointment.currentRow()
        if row_index >= 0:
            appointment_id = self.ui.tableWidgetAppointment.item(row_index, 0).text()  # Получаем ID записи
            patient_row = self.ui.tableWidgetPatients.currentRow()  # Получаем выбранного пациента
            doctor_row = self.ui.tableWidgetDoc.currentRow()  # Получаем выбранного врача

            # Получаем значения из таблицы для обновления
            appointment_date = self.ui.tableWidgetAppointment.item(row_index, 3).text()  # Дата приема
            date_of_application = self.ui.tableWidgetAppointment.item(row_index, 4).text()  # Дата обращения
            status = self.ui.tableWidgetAppointment.item(row_index, 5).text()  # Статус посещения

            if patient_row >= 0 and doctor_row >= 0:
                patient_id = self.ui.tableWidgetPatients.item(patient_row, 0).text()  # ID пациента
                doctor_id = self.ui.tableWidgetDoc.item(doctor_row, 0).text()  # ID врача

                connection = sqlite3.connect('db/hospital_management.db')
                cursor = connection.cursor()
                cursor.execute("""
                    UPDATE Appointments 
                    SET patient_id = ?, doctor_id = ?, appointment_date = ?, date_of_application = ?, status = ?
                    WHERE id = ?
                """, (patient_id, doctor_id, appointment_date, date_of_application, status, appointment_id))
                connection.commit()
                connection.close()

                self.load_appointments()  # Обновляем таблицу записей
            else:
                QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите пациента и врача для редактирования записи.")
        else:
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите запись для редактирования.")

    def load_appointments(self):
        connection = sqlite3.connect('db/hospital_management.db')
        cursor = connection.cursor()
        cursor.execute(
            "SELECT Appointments.id, Patients.full_name, Doctors.full_name, Appointments.appointment_date, Appointments.date_of_application, Appointments.status "
            "FROM Appointments "
            "JOIN Patients ON Appointments.patient_id = Patients.id "
            "JOIN Doctors ON Appointments.doctor_id = Doctors.id"
        )
        self.ui.tableWidgetAppointment.setRowCount(0)
        for row in cursor.fetchall():
            row_position = self.ui.tableWidgetAppointment.rowCount()
            self.ui.tableWidgetAppointment.insertRow(row_position)
            for column, data in enumerate(row):
                self.ui.tableWidgetAppointment.setItem(row_position, column, QtWidgets.QTableWidgetItem(str(data)))
        connection.close()
        self.save_appointments_to_word()  # Если нужно будет вызывать при загрузке или по нажатию кнопки

    def save_appointments(self):
        """Сохраняет все записи из таблицы в базу данных."""
        connection = None
        try:
            connection = sqlite3.connect('db/hospital_management.db')
            cursor = connection.cursor()

            for row in range(self.ui.tableWidgetAppointment.rowCount()):
                appointment_id = self.ui.tableWidgetAppointment.item(row, 0).text()
                patient_name = self.ui.tableWidgetAppointment.item(row, 1).text()
                doctor_name = self.ui.tableWidgetAppointment.item(row, 2).text()
                appointment_date = self.ui.tableWidgetAppointment.item(row, 3).text()
                date_of_application = self.ui.tableWidgetAppointment.item(row, 4).text()
                status = self.ui.tableWidgetAppointment.item(row, 5).text()

                patient_id = self.get_patient_id(patient_name) if hasattr(self, 'get_patient_id') else None
                doctor_id = self.get_doctor_id(doctor_name) if hasattr(self, 'get_doctor_id') else None

                if patient_id and doctor_id and appointment_date and date_of_application and status:
                    # Проверяем, существует ли уже запись с таким же пациентом, врачом и датами
                    cursor.execute("""
                        SELECT COUNT(*) FROM Appointments 
                        WHERE patient_id = ? AND doctor_id = ? 
                        AND appointment_date = ? AND date_of_application = ? AND status = ?
                    """, (patient_id, doctor_id, appointment_date, date_of_application, status))
                    count = cursor.fetchone()[0]

                    if count == 0:  # Если записей нет, добавляем новую
                        cursor.execute("""
                                    INSERT INTO Appointments (patient_id, doctor_id, appointment_date, date_of_application, status) 
                                    VALUES (?, ?, ?, ?, ?)
                                """, (patient_id, doctor_id, appointment_date, date_of_application, status))
                    else:
                        QMessageBox.warning(self, "Ошибка",
                                            f"Запись для пациента '{patient_name}' и врача '{doctor_name}' на данную дату уже существует.")
                else:
                    QMessageBox.warning(self, "Ошибка",
                                        f"Запись {row + 1} содержит пустые поля. Пожалуйста, проверьте данные.")

            connection.commit()
            self.load_appointments()
            self.load_patient_statistics()
            QMessageBox.information(self, "Информация",
                                    "Данные о надежности пациентов обновлены.")
        except sqlite3.Error as e:
            QMessageBox.critical(self, "Ошибка базы данных", f"Произошла ошибка при сохранении: {e}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Произошла непредвиденная ошибка: {e}")
        finally:
            if connection:
                connection.close()

    def get_patient_id(self, patient_name):
        connection = sqlite3.connect('db/hospital_management.db')
        cursor = connection.cursor()
        # Изменяем 'name' на 'full_name'
        cursor.execute("SELECT id FROM Patients WHERE full_name = ?", (patient_name,))
        result = cursor.fetchone()
        connection.close()
        return result[0] if result else None

    def get_doctor_id(self, doctor_name):
        connection = sqlite3.connect('db/hospital_management.db')
        cursor = connection.cursor()
        # Изменяем 'name' на 'full_name'
        cursor.execute("SELECT id FROM Doctors WHERE full_name = ?", (doctor_name,))
        result = cursor.fetchone()
        connection.close()
        return result[0] if result else None

    def calculate_patient_reliability(self):
        """Выводит надежность пациентов на основе посещаемости."""
        connection = sqlite3.connect('db/hospital_management.db')
        cursor = connection.cursor()

        # SQL-запрос для подсчета посещаемости пациентов с JOIN для получения ФИО
        cursor.execute("""
            SELECT Patients.full_name, 
                   COUNT(*) AS total_appointments, 
                   SUM(CASE WHEN status = 'Пришел' THEN 1 ELSE 0 END) AS attended 
            FROM Appointments 
            JOIN Patients ON Appointments.patient_id = Patients.id
            GROUP BY Patients.full_name
        """)

        reliability_data = []
        for row in cursor.fetchall():
            patient_full_name = row[0]
            total_appointments = row[1]
            attended = row[2]

            # Вычисляем надежность
            reliability = (attended / total_appointments * 100) if total_appointments > 0 else 0
            reliability_data.append((patient_full_name, reliability))  # Сохраняем ФИО пациента

        connection.close()
        return reliability_data

    def load_patient_statistics(self):
        """Загружает данные о надежности пациентов и отображает в текстовом формате на странице статистики."""
        reliability_data = self.calculate_patient_reliability()

        # Создаем строку для отображения данных о надежности
        reliability_text = "<b>Надежность пациентов:</b><br>"  # Используем <br> для переноса строки

        for full_name, reliability in reliability_data:
            # Форматируем строки, чтобы ФИО занимало 30 символов, а процент 10 символов,
            # и добавляем жирный шрифт для процента
            reliability_text += f"{full_name:<30}: {reliability:.2f}%<br>"

        self.ui.StatisticName.setText(reliability_text)  # Устанавливаем текст для QLabel Status
        self.ui.StatisticName.setTextFormat(QtCore.Qt.RichText)  # Устанавливаем формат текста как RichText

    def update_patient_data(self, row, column):
        """Обновляет данные пациента в базе данных при изменении значений в таблице."""
        patient_id = self.ui.tableWidgetPatients.item(row, 0).text()  # Получаем ID пациента
        full_name = self.ui.tableWidgetPatients.item(row, 1).text()
        birthday = self.ui.tableWidgetPatients.item(row, 2).text()
        address = self.ui.tableWidgetPatients.item(row, 3).text()
        gender = self.ui.tableWidgetPatients.item(row, 4).text()
        phone_number = self.ui.tableWidgetPatients.item(row, 5).text()

        connection = sqlite3.connect('db/hospital_management.db')
        cursor = connection.cursor()
        cursor.execute("""
            UPDATE Patients 
            SET full_name = ?, birthday = ?, address = ?, gender = ?, phone_number = ?
            WHERE id = ?
        """, (full_name, birthday, address, gender, phone_number, patient_id))
        connection.commit()
        connection.close()

    def update_doctor_data(self, row, column):
        doctor_id = self.ui.tableWidgetDoc.item(row, 0).text()  # Получаем ID врача
        full_name = self.ui.tableWidgetDoc.item(row, 1).text()
        job_title = self.ui.tableWidgetDoc.item(row, 2).text()

        connection = sqlite3.connect('db/hospital_management.db')
        cursor = connection.cursor()
        cursor.execute("""
            UPDATE Doctors 
            SET full_name = ?, job_title = ?
            WHERE id = ?
        """, (full_name, job_title, doctor_id))
        connection.commit()
        connection.close()

    def save_appointments_to_word(self):
        """Сохраняет записи из таблицы в файл Word."""
        document = Document()
        document.add_heading('Записи пациентов', level=1)

        # Добавляем таблицу
        table = document.add_table(rows=1, cols=3)  # 1 заголовок и 3 колонки
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ФИО пациента'
        hdr_cells[1].text = 'ФИО врача'
        hdr_cells[2].text = 'Время записи'

        # Заполняем таблицу данными
        connection = sqlite3.connect('db/hospital_management.db')
        cursor = connection.cursor()
        cursor.execute("""
            SELECT Patients.full_name, Doctors.full_name, Appointments.appointment_date
            FROM Appointments
            JOIN Patients ON Appointments.patient_id = Patients.id
            JOIN Doctors ON Appointments.doctor_id = Doctors.id
        """)
        for patient_name, doctor_name, appointment_date in cursor.fetchall():
            row_cells = table.add_row().cells
            row_cells[0].text = patient_name
            row_cells[1].text = doctor_name
            row_cells[2].text = appointment_date

        connection.close()
        # Указываем путь для сохранения документа
        save_folder = os.path.join(os.path.dirname(__file__), 'static', 'files')  # Путь к папке
        os.makedirs(save_folder, exist_ok=True)  # Создание папки, если она не существует
        document_path = os.path.join(save_folder, 'appointments.docx')  # Полный путь к файлу

        # Сохранение документа
        document.save(document_path)

    def highlight_patient(self):
        try:
            patient_id = int(self.ui.searchLineEdit.text().strip())  # Получаем ID из поля ввода
            row_count = self.ui.tableWidgetPatients.rowCount()

            found_patient = False  # Переменная для отслеживания найденных пациентов

            # Сброс предыдущей подсветки
            for row in range(row_count):
                for column in range(self.ui.tableWidgetPatients.columnCount()):
                    item = self.ui.tableWidgetPatients.item(row, column)
                    if item is not None:
                        item.setBackground(QtGui.QColor(255, 255, 255))  # Сбрасываем цвет фона на белый

            # Подсвечиваем пациента в таблице
            for row in range(row_count):
                item = self.ui.tableWidgetPatients.item(row, 0)
                if item is not None and item.text() == str(patient_id):  # Проверяем ID
                    found_patient = True
                    # Изменяем цвет фона строки
                    for column in range(self.ui.tableWidgetPatients.columnCount()):
                        item = self.ui.tableWidgetPatients.item(row, column)
                        if item is not None:
                            item.setBackground(QtGui.QColor(240, 240, 240))  # Серый цвет для подсветки

            if not found_patient:
                QtWidgets.QMessageBox.warning(self, "Ошибка", "Пациент не найден.")
        except ValueError:
            QtWidgets.QMessageBox.warning(self, "Ошибка", "Введите корректный ID пациента.")
        except Exception as e:
            QtWidgets.QMessageBox.warning(self, "Ошибка", f"Произошла ошибка: {str(e)}")

    def get_patient_by_id(self, patient_id):
        """Получение информации о пациенте по ID."""
        connection = sqlite3.connect('db/hospital_management.db')
        cursor = connection.cursor()
        cursor.execute("SELECT * FROM Patients WHERE id = ?", (patient_id,))
        patient_data = cursor.fetchone()  # Получаем одну запись

        connection.close()

        if patient_data:
            return patient_data  # Возвращаем данные о пациенте
        else:
            return None  # Если пациент не найден, возвращаем None

    def show_right_menu(self):
        """Показать правое меню"""
        self.ui.RightMenu.setMaximumSize(QtCore.QSize(300, 16777215))  # Задаём максимальные размеры для видимости
        self.ui.RightMenu.setVisible(True)  # Делаем правое меню видимым
        self.ui.Edit.setVisible(False)  # Скрываем кнопку Edit
        self.ui.Delete.setVisible(False)  # Скрываем кнопку удаления при открытии правого меню
        self.ui.Edit2.setVisible(False)
        self.ui.Delete2.setVisible(False)
        self.ui.Edit3.setVisible(False)
        self.ui.Delete3.setVisible(False)
        self.ui.RightMenuDoctors.setMaximumSize(QtCore.QSize(0, 16777215))  # Скрываем меню врачей
        self.ui.RightMenuDoctors.setVisible(False)  # Скрываем панель для докторов

    def show_right_menu_doctors(self):
        """Показать правое меню для врачей."""
        self.ui.RightMenuDoctors.setMaximumSize(QtCore.QSize(300, 16777215))
        self.ui.RightMenuDoctors.setVisible(True)  # Делаем видимым меню врачей
        self.ui.Edit.setVisible(False)  # Скрываем кнопку Edit
        self.ui.Delete.setVisible(False)  # Скрываем кнопку удаления при открытии правого меню
        self.ui.Edit2.setVisible(False)
        self.ui.Delete2.setVisible(False)
        self.ui.Edit3.setVisible(False)
        self.ui.Delete3.setVisible(False)
        self.ui.RightMenu.setMaximumSize(QtCore.QSize(0, 16777215))  # Скрываем меню пациентов
        self.ui.RightMenu.setVisible(False)  # Скрываем правое меню для пациентов

    def hide_right_menu_doctors(self):
        """Скрыть правое меню для врачей."""
        self.ui.RightMenuDoctors.setMaximumSize(
            QtCore.QSize(0, 16777215))  # Устанавливаем максимальные размеры для скрытия
        self.ui.RightMenuDoctors.setVisible(False)  # Скрываем правое меню для врачей
        self.ui.Patients.setVisible(True)  # Отображаем кнопку "Пациенты"
        self.ui.Doctors.setVisible(True)  # Отображаем кнопку "Доктор", если необходимо
        self.ui.Edit.setVisible(True)  # Показываем кнопку Edit
        self.ui.Delete.setVisible(True)
        self.ui.Edit2.setVisible(True)
        self.ui.Delete2.setVisible(True)
        self.ui.Edit3.setVisible(True)
        self.ui.Delete3.setVisible(True)

    def hide_right_menu(self):
        """Скрыть правое меню"""
        self.ui.RightMenu.setMaximumSize(QtCore.QSize(0, 16777215))  # Задаём максимальные размеры для скрытия
        self.ui.RightMenu.setVisible(False)  # Скрываем правое меню
        self.ui.Patients.setVisible(True)  # Отображаем кнопку "Пациенты"
        self.ui.Doctors.setVisible(True)  # Отображаем кнопку "Доктор", если необходимо
        self.ui.Edit.setVisible(True)  # Показываем кнопку Edit
        self.ui.Delete.setVisible(True)
        self.ui.Edit2.setVisible(True)
        self.ui.Delete2.setVisible(True)
        self.ui.Edit3.setVisible(True)
        self.ui.Delete3.setVisible(True)

    def exit_to_login(self):
        """Скрывает все окна и показывает окно логина."""
        self.hide()  # Скрываем главное окно
        login_widget = Login()  # Создаем экземпляр окна логина
        if login_widget.exec_() == QtWidgets.QDialog.Accepted:
            self.show()  # Показываем основное окно снова, если логин был успешным

    def hideAllPages(self):
        """Скрывает все страницы."""
        self.ui.Labpage.setVisible(False)
        self.ui.ReportpagePatients.setVisible(False)
        self.ui.Statisticpage.setVisible(False)
        self.ui.Settingspage.setVisible(False)
        self.ui.Helppage.setVisible(False)
        self.ui.ReportpageDoc.setVisible(False)
        self.ui.ReportpageAppointment.setVisible(False)

    def showWidget(self, page):
        """Показывает нужный виджет и скрывает все остальные."""
        self.hideAllPages()  # Скрыть все страницы

        # Показываем соответствующий виджет
        if page == 'lab':
            self.ui.Labpage.setVisible(True)
        elif page == 'report':
            self.ui.ReportpagePatients.setVisible(True)
        elif page == 'statistics':
            self.ui.Statisticpage.setVisible(True)
        elif page == 'settings':
            self.ui.Settingspage.setVisible(True)
        elif page == 'help':
            self.ui.Helppage.setVisible(True)

    def hideShowLeftBar(self):
        """Скрывает или показывает левую панель меню."""
        isVisible = self.ui.LeftMenu.isVisible()  # Проверка текущей видимости
        self.ui.LeftMenu.setVisible(not isVisible)  # Переключение видимости

    def show_patients_table(self):
        """Показать таблицу пациентов."""
        self.hideAllPages()  # Скрываем все остальные страницы
        self.ui.ReportpagePatients.setVisible(True)  # Показываем страницу пациентов
        self.ui.Edit.setVisible(True)  # Показываем кнопку Edit
        self.ui.Delete.setVisible(True)
        self.ui.stackedWidget.setCurrentWidget(self.ui.ReportpagePatients)  # Устанавливаем текущий виджет

    def show_doctors_table(self):
        """Показать таблицу врачей."""
        self.hideAllPages()  # Скрываем все остальные страницы
        self.ui.ReportpageDoc.setVisible(True)  # Показываем страницу врачей
        self.ui.Edit2.setVisible(True)  # Показываем кнопку Edit
        self.ui.Delete2.setVisible(True)
        self.ui.stackedWidget.setCurrentWidget(self.ui.ReportpageDoc)  # Устанавливаем текущий виджет

    def show_appointments_table(self):
        """Показать таблицу записей на прием."""
        self.hideAllPages()  # Скрываем все остальные страницы
        self.ui.ReportpageAppointment.setVisible(True)  # Показываем страницу записей
        self.ui.Edit3.setVisible(True)  # Показываем кнопку Edit
        self.ui.Delete3.setVisible(True)
        self.ui.stackedWidget.setCurrentWidget(self.ui.ReportpageAppointment)  # Устанавливаем текущий виджет


if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)

    # Показываем окно логина
    login_widget = Login()
    if login_widget.exec_() == QtWidgets.QDialog.Accepted:  # Показываем окно логина и ждем результата
        main_window = MyApp()
        main_window.show()  # Показываем основное окно
        sys.exit(app.exec_())

